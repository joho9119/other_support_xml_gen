import re
import sys
import io
from pathlib import Path


import requests
from docx import Document
from docx.document import Document as _Document

CMD_ARGS = sys.argv
OTHER_SUPPORT_BLANK = "https://grants.nih.gov/sites/default/files/other-support-format-page-rev-10-2021.docx"
OTHER_SUPPORT_SAMPLE = "https://grants.nih.gov/sites/default/files/other-support-sample-7-20-2021.docx"
FIELD_LABELS = {
    "section_header": re.compile(r"^(ACTIVE|PENDING|IN-KIND)", re.IGNORECASE),
    "name_id": re.compile(r"Name of Individual:\s*(.+?)(?:\s+Commons ID:.*)?$", re.IGNORECASE),
    "project_title": re.compile(r"Title:\s*", re.IGNORECASE),
    "major_goals": re.compile(r"Major Goals:\s*", re.IGNORECASE),
    "status": re.compile(r"Status of Support:\s*", re.IGNORECASE),
    "role": re.compile(r"Role:\s*", re.IGNORECASE),
    "project_number": re.compile(r"Project Number:\s*", re.IGNORECASE),
    "pd_pi": re.compile(r"Name of PD/PI:\s*", re.IGNORECASE),
    "source": re.compile(r"Source of Support:\s*", re.IGNORECASE),
    "place": re.compile(r"(?:Primary )?Place of Performance:\s*", re.IGNORECASE),
    "dates": re.compile(r"Project.*?Date.*?:", re.IGNORECASE),
    "amount": re.compile(r"Total Award Amount.*?:", re.IGNORECASE),
    "overlap": re.compile(r"\*?Overlap\s*:\s*", re.IGNORECASE),
    "person_months_stopper": re.compile(r"Person\s*Months", re.IGNORECASE)
}

DATE_EXTRACTOR = re.compile(r"(\d{1,2}/\d{1,2}/\d{2,4}|\d{1,2}/\d{2,4})\s*[-–]\s*(\d{1,2}/\d{1,2}/\d{2,4}|\d{1,2}/\d{2,4})")
AMOUNT_EXTRACTOR = re.compile(r"\$?([\d,]+)")



def load_document(doc_input) -> _Document:
    if str(doc_input).startswith("http"):
        print(f"Fetching document from URL: {doc_input}")
        response = requests.get(doc_input)
        response.raise_for_status()
        doc_source = io.BytesIO(response.content)
    else:
        if not str(doc_input).lower().endswith(".docx"):
            raise TypeError("Provided Word doc must end in .docx")
        doc_path = Path(doc_input)
        doc_path.is_file()
        if not doc_path.is_file():
            if not doc_path.exists():
                raise ValueError(f"Provided input does not point to a valid file path. (Input: {doc_input}")
            if doc_path.is_dir():
                raise TypeError(f"Provided input must point to a file. (Pointing to folder: {doc_path.name}).")
        doc_source = doc_path

    doc = Document(doc_source)
    print(f"Successfully loaded document. Paragraphs: {len(doc.paragraphs)}")
    return doc


def parse_other_support(doc_input):
    word_doc = load_document(doc_input)
    print(f"File loaded: {word_doc.paragraphs}")


def main():
    if len(CMD_ARGS) < 2:
        print("Usage: python parse_os.py {path to OS document}")
        return
    else:
        if CMD_ARGS[1:][0].lower() == "sample":
            doc_input = OTHER_SUPPORT_SAMPLE
        else:
            doc_input = CMD_ARGS[1:][0]
    parse_other_support(doc_input)


if __name__ == "__main__":
    main()
